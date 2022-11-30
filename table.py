import xlsxwriter
from docx import Document
from docx.shared import Inches
from random import *
guestnames = ['Ернар', 'Амир', 'Даурен', 'Арман', 'Алихан', 'Куттубек', 'Даниал']

cont = 'hello'
while cont != 'Excel' and cont != 'Word':
 cont = input('Введите желаемый формат документа (Word или Excel):')

if cont == 'Word':
 for x in guestnames:
  document = Document()
  document.add_heading('Приглашение на конкурс проектов', 0)
  p = document.add_paragraph(x + ', спешим пригласить Вашу команда на это мероприятие')
  p.add_run(' PizzaPitch').bold = True
  p.add_run(', которая состоится 18.11.2022 в Toraighyrov University.')
  p.add_run('\nСпешу вам сообщить, что уже совсем скоро вы будете демострировать свой проект.')
  p.add_run('\nТак же по проекту, вам необходимо оформить вашу идею как бизнес/стартап проект. Составить презентацию и  при необходимости вышлю вам план презентации')
  document.add_heading('Мы ждём вас, на нашем мероприятии! \nВаш регистрационный номер: '+ str(randint(0,1000)), level=1)
  document.add_paragraph('')
  document.add_picture('1.jpeg', width=Inches(5.25))
  document.add_page_break()
  document.save(x + '.docx')
  print('Выполнено')
else:
 for x in guestnames:
  my_file = (x + '.xlsx')
  book = xlsxwriter.Workbook(my_file)
  sheet = book.add_worksheet()
  sheet.set_column('B:B', 80)
  bold = book.add_format({'bold': True})
  sheet.write('B1', 'Приглашение на конкурс проектов', bold)
  sheet.write('B2', x + ', спешим пригласить Вашу команда на это мероприятие PizzaPitch.')
  sheet.write('B3', 'Которая состоится 18.11.2022 в Toraighyrov University.')
  sheet.write('B4', 'Спешу вам сообщить, что уже совсем скоро вы будете демострировать свой проект.')
  sheet.write('B5', 'Так же по проекту, вам необходимо оформить вашу идею как бизнес/стартап проект.')
  sheet.write('B6', 'Составить презентацию и  при необходимости вышлю вам план презентации')
  sheet.write('B6', 'Мы ждём вас, на нашем мероприятии!', bold)
  sheet.write('B6', 'Ваш регистрационный номер: ' + str(randint(0,1000)), bold)
  sheet.insert_image('B7', '1.jpeg', {'x_scale': 0.50, 'y_scale': 0.50})
  book.close()
  print('Выполнено')
