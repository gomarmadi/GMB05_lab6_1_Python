from docx import Document
from docx.shared import Inches
from random import *
guestnames = ['Ернар', 'Амир', 'Даурен', 'Арман', 'Алихан', 'Куттубек', 'Даниал']
for x in guestnames:
 document = Document()
 document.add_heading('Приглашение на конкурс проектов', 0)
 p = document.add_paragraph(x + ', спешим пригласить Вашу команда на это мероприятие')
 p.add_run(' PizzaPitch').bold = True
 p.add_run(', которая состоится 18.11.2022 в Toraighyrov University.')
 p.add_run('\nСпешу вам сообщить, что уже совсем скоро вы будете демострировать свой проект.')
 p.add_run('\nТак же по проекту, вам необходимо оформить вашу идею как бизнес/стартап проект. Составить презентацию и  при необходимости вышлю вам план презентации')
 document.add_heading('Мы ждём вас, на нашем мероприятии! \nВаш регистрационный номер: '+ str(randint(0,1000)), level=1)
 document.add_paragraph('')
 document.add_picture('1.jpeg', width=Inches(5.25))
 document.add_page_break()
 document.save(x + '.docx')